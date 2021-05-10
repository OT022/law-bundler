import argparse
import csv
import fs
import logging
import os
import sys
import datetime
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Cm
from docx2pdf import convert
from PyPDF2 import PdfFileMerger, PdfFileReader, PdfFileWriter
from reportlab.lib.units import mm
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

from batch_to_vol_helper import split_to_volumes

pdfmetrics.registerFont(TTFont("Calibri", "Calibri.ttf"))


def getInputFromExportedCSV(fname: str, inputdir: str):
    """ 
    - takes a file name and opens it
    - reads it to a list and returns that list of data
    input_data - list holding the data read from the file
    fname - name of csv file being read from
    inputdir - directory name of the input files
    """
    input_data = []
    with open("{}/{}.csv".format(inputdir, fname), "r", encoding="utf-8") as export:
        reader = csv.reader(export)
        for r in reader:
            input_data.append(r)
    return input_data


def getPageIndex(data_list: list, inputdir: str):
    """ 
    - takes data_list and loops through it
    - gets the number of pages in each pdf
    - appends the page index to the end of each list in the nested list data_list
    data_list - the nested list of data from the input file
    curr_pdf - is the current pdf being opened in the loop
    fname - the name of each pdf in the bundle
    """
    page_index = 1
    bundle_counter = 0

    for lines in data_list:
        if lines[0] != "Control Number":

            if lines[1] == "":
                fname = "{}/{}.pdf".format(inputdir, lines[0])
            else:
                fname = "{}/{}.pdf".format(inputdir, lines[1])

            f = open(fname, "rb")
            curr_pdf = PdfFileReader(f)

            num_pages = curr_pdf.getNumPages()
            data_list[bundle_counter].append(str(page_index))
            data_list[bundle_counter].append(str(num_pages))
            page_index += num_pages

        bundle_counter += 1

    return data_list


def createIndexPage(index_data: list, outputdir: str, filename="index.docx"):
    """
    - taking the index data and populating a table in the index template word doc
    - saving that doc as index.docx
    index_data - the list of index page data including comntrol number, data and contents
    table - is the table in the template doc
    outputdir - is the directory of the output data 
    """

    # creating a word document
    word_doc = Document("table_template.docx")

    # * used to indicate document included at external request
    asterisk = False

    # customizing the table
    # table = word_doc.add_table(rows=len(export_data), cols=4)
    # table.style= "Index2"
    table = word_doc.tables[0]

    # adding the title row to the index page table
    table.cell(0, 0).text = "Inquiry Reference Number"
    table.cell(0, 1).text = "Date"
    table.cell(0, 2).text = "Description"
    table.cell(0, 3).text = "Page"

    for counter in range(0, len(index_data)):
        table.add_row()
        # checking if witness document id is present otherise using control number
        if index_data[counter][1] == "":
            if index_data[counter][0][0] == "S":
                # counter + 1 to account for header in .docx table rows
                table.rows[counter + 1].cells[0].text = index_data[counter][0] # + "*"
                asterisk = True
            else:
                table.rows[counter + 1].cells[0].text = index_data[counter][0]
        else:
            if index_data[counter][0] != "" and index_data[counter][0][0] == "S":
                table.rows[counter + 1].cells[0].text = index_data[counter][1] # + "*"
                asterisk = True
            else:
                table.rows[counter + 1].cells[0].text = index_data[counter][1]

        # if undated, display in index
        if index_data[counter][2] != "" and index_data[counter][4] == "Yes":            
             table.rows[counter + 1].cells[1].text = index_data[counter][2] + " (estimated)"
        elif index_data[counter][4] == "Undated":
            table.rows[counter + 1].cells[1].text = "Undated"
        elif index_data[counter][2] != "":
            table.rows[counter + 1].cells[1].text = index_data[counter][2]
        else:
            table.rows[counter + 1].cells[1].text = index_data[counter][6]

        # add content description
        table.rows[counter + 1].cells[2].text = index_data[counter][3]

        # add page number and center style
        cell = table.rows[counter + 1].cells[3]
        cell.text = index_data[counter][7]  # need to get index_data_page
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    ## ignore

    asterisk = False
    if asterisk:
        para = word_doc.add_paragraph(
            "*Document included at request of TSA"
        )  # TODO make the requestor a variable that can be changed
    # save word doc as index.docx
    word_doc.save("{}/{}".format(outputdir, filename))
    convert("{}/{}".format(outputdir, filename))
    file_docx = "{}/{}".format(outputdir, filename)
    return file_docx.replace("docx", "pdf")


def createCoverPage(bundle_title, vol_no
    , cover_page_date, is_statement, is_draft, outputdir
):
    """
    - takes the cover_page_title/date and edits the word doc template 
    - saves it as cover.docx
    cover_page_title - str that populates the title of the cover page
    cover_page_data - str of the date on the cover page
    """
    cover_page = Document("Cover Page Template.docx")
    cover_page.paragraphs[3].text = "FAITH BASED HEARINGS (PHASE II)"  # to_upper_case
    styles = cover_page.paragraphs[3].style
    styles.font.bold = True
    
    if is_statement:
        cover_page.paragraphs[4].text = "WITNESS STATEMENT OF"
        cover_page.paragraphs[5].text = bundle_title.upper()
        cover_page.paragraphs[6].text = ""

    if not is_statement:
        cover_page.paragraphs[4].text = "BUNDLE OF EVIDENCE"
        cover_page.paragraphs[5].text = bundle_title.upper()
        cover_page.paragraphs[6].text = "Volume {}".format(vol_no).upper()
    
    if is_draft:
        cover_page.paragraphs[4].text = "DRAFT AS AT {}".format(
            datetime.datetime.now().isoformat()
        )
        cover_page.paragraphs[5].text = "DRAFT"
        cover_page.paragraphs[6].text = "DRAFT"
    cover_page.paragraphs[7].text = cover_page_date  # to_upper_case

    return cover_page


def mergePdfsIntoBundle(
    cover_page_path: str,
    index_pdf_path: str,
    exhibit_list: list,
    statement_list: list,
    is_statement: bool,
    inputdir: str,
    outputdir: str,
    outputname: str,
    name: str,
):
    """
    - appends index and cover pdfs to merger
    - loops through pdfs and appends them to merger
    - then merges the pdf and writes it to outputname in outputdir
    merger - is the PyPDF2 Class that merges all the pdfs in the list
    outputdir - is the directory of the output file
    """
    merger = PdfFileMerger()

    print("statement_list", statement_list)

    # bundle is a joint bundle - statement + exhibits
    if not is_statement:
        merger.append(cover_page_path)
        for statement in statement_list:
            # checking if it is first list in data_list
            if statement[0] != "Control Number":
                # if statement[1] == "":
                fname = "{}/{}.pdf".format(inputdir, statement[0])
                # else:
                    # fname = "{}/{}.pdf".format(inputdir, statement[1])

                curr_pdf = PdfFileReader(open(fname, "rb"))

            merger.append(curr_pdf)
        merger.append(index_pdf_path)
        for exhibit in exhibit_list:
            # checking if it is first list in data_list
            if exhibit[0] != "Control Number":
                # find file by ctrl number or witness doc id
                if exhibit[1] == "":
                    fname = "{}/{}.pdf".format(inputdir, exhibit[0])
                    logging.info(name + ":" + exhibit[0] + " added to volume")

                else:
                    fname = "{}/{}.pdf".format(inputdir, exhibit[1])
                    logging.info(name + ":" + exhibit[1] + " added to volume")

                curr_pdf = PdfFileReader(open(fname, "rb"))

                merger.append(curr_pdf)

    # bundle is a statement bundle only
    else:
        merger.append(cover_page_path)
        for statement in statement_list:
            # checking if it is first list in data_list
            if statement[0] != "Control Number":
                # find file by ctrl number or witness doc id
                if statement[1] == "":
                    fname = "{}/{}.pdf".format(inputdir, statement[0])
                else:
                    fname = "{}/{}.pdf".format(inputdir, statement[1])

                curr_pdf = PdfFileReader(open(fname, "rb"))

                merger.append(curr_pdf)

    while True:
        try:
            merger.write("{}/{}".format(outputdir, outputname))
            break
        except PermissionError:
            print("\n" + outputname + " is still open, close it to continue")
            input("press enter to continue:")

    # we've written pdf,
    # now generate pagination overlay and merge again
    if not is_statement:
        paginateVolume(
            "{}/{}".format(outputdir,outputname),
            index_pdf_path,
            outputdir,
            outputname,
            name,
            int(exhibit_list[0][7])
            )


def createPagePdf(bundle_length: int, index_page_length: int, tmp: str, start_pagination_at: int):
    """
    creates a pdf with page numbers and overlays it on top 
    of the exihits and statements
    start_pagination_at - is the page after which the pagination should start
    """
    c = canvas.Canvas(tmp)
    for i in range(0, bundle_length):
        if i > index_page_length:
            page_num = start_pagination_at + i - index_page_length - 1 # -1 to account for cover page?
            print(i, page_num)
            c.setFont("Calibri", 10)
            c.drawString(195 * mm, (20) * mm, str(page_num))
            c.showPage()
        else:
            c.showPage()
    c.save()


def paginateVolume(
    path_to_merged_vol: str,
    path_to_index: str,
    outputdir: str,
    outputname: str,
    bundle_title: str,
    index_start: int
    ):

    
    index_len = 0

    # get index_len
    with open(path_to_index, "rb") as idx_tmp:
        index_len = PdfFileReader(idx_tmp).getNumPages()

    # create dummy pdf holder
    tmp = "{}/__tmp.pdf".format(outputdir)

    with open(path_to_merged_vol, "rb") as vol_tmp:
        vol_pdf = PdfFileReader(vol_tmp, strict=True)
        vol_len = vol_pdf.getNumPages()

        output = PdfFileWriter()

        # create new PDF with page numbers
        createPagePdf(vol_len, index_len, tmp, index_start)

        newpath = "{}/{}_paginated.pdf".format(outputdir, bundle_title.replace(" ", "_"))

        with open(tmp, "rb") as num_pdf_tmp:
            numberPdf = PdfFileReader(num_pdf_tmp)
            # iterate pages
            for p in range(vol_len):
                page = vol_pdf.getPage(p)
                numberLayer = numberPdf.getPage(p)
                # merge number page with actual page
                page.mergePage(numberLayer)
                output.addPage(page)
                print("merge page ", p)

            while True:
                try:
                    with open(newpath, "wb") as f:
                        output.write(f)
                        logging.info(
                            bundle_title + ":paginated"
                        )
                    break
                except PermissionError:
                    print("\n" + bundle_title + " - STATEMENT AND EXHIBITS.pdf")
                    input("press enter to continue")
        os.remove(tmp)


def main():

    # data of all bundles - comes from order.csv
    master_data = []
    with open("src/order.csv", "r") as order:
        reader = csv.reader(order)
        for r in reader:
            master_data.append(r)

    print("master data", master_data)

    for i in range(len(master_data)):
        bundle = master_data[i]

        if i == 0:
            print("row headers are", master_data[i])
            continue

        bundle_title = bundle[0]
        bundle_date = bundle[1]

        should_skip = bool(int(bundle[2]))
        if not should_skip:
            is_statement = bool(int(bundle[4]))
            is_draft = bool(int(bundle[3]))  # pull is_draft bool val
            inputfile = bundle_title
            inputdir = "src/{}".format(inputfile)
            outputname = bundle_title + " - STATEMENT.pdf"
            outputdir = "output/{}".format(inputfile)
            input_data_file = "export"

            # create output dir for bundle if it doesn't exist
            if not os.path.isdir("output/{}".format(bundle_title)):
                os.mkdir("output/{}".format(bundle_title))

            # create log file
            logging.basicConfig(
                filename="{}/{}.log".format(outputdir, bundle_title),
                format="%(asctime)s:%(message)s",
                datefmt="%Y-%m-%dT%H:%M:%S%z",
                level=logging.INFO,
            )

            logging.info(bundle_title + ":started")
            export_data = getInputFromExportedCSV(input_data_file, inputdir)

            # cover page
            cover_page_date = bundle_date.upper()

            # populate index data
            export_data = getPageIndex(export_data, inputdir)

            # organise into volumes
            # TODO  first volume is statement only
            volumes = split_to_volumes(export_data, 1000)

            for i in range(len(volumes)):
                vol_no = i + 1
                volume_title = "{} VOLUME {}".format(bundle_title, vol_no)

                cover_page = createCoverPage(bundle_title, vol_no, cover_page_date, is_statement, is_draft, outputdir
                )

                cover_page_file_name = "{}/cover-vol-{}.docx".format(outputdir, vol_no)
                index_pdf_path = createIndexPage(
                    volumes[i], outputdir, "index-vol-{}.docx".format(vol_no)
                )

                cover_page_pdf_path = "{}/cover-vol-{}.pdf".format(outputdir, vol_no)
                output_file_name = "{}.pdf".format(volume_title.replace(" ", "_"))
                cover_page.save(cover_page_file_name)
                convert(cover_page_file_name)

                mergePdfsIntoBundle(
                    cover_page_pdf_path,
                    index_pdf_path,
                    volumes[i],
                    [],  # ignore statements for now - all are statements - just pass in
                    False,  # as above
                    inputdir,
                    outputdir,
                    output_file_name,
                    volume_title,
                )
                logging.info(output_file_name + ":created")


if __name__ == "__main__":
    main()
